"""
简历格式化工具
- 上传 TXT 文件
- AI 美化内容
- 导出 DOC 文件
"""

import os
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

load_dotenv()

# 兼容本地和云端部署
api_key = os.getenv("DASHSCOPE_API_KEY")

client = OpenAI(
    api_key=api_key,
    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
)

st.set_page_config(page_title="简历格式化工具", page_icon="📄", layout="wide")

st.title("📄 简历格式化工具")
st.write("上传 TXT 简历，AI 帮你美化并导出为 DOC 格式")

# 上传文件
uploaded_file = st.file_uploader("上传简历 (TXT 格式)", type=["txt"])

# 或直接输入
resume_text = st.text_area(
    "或直接粘贴简历内容",
    placeholder="姓名：张三\n联系方式：138xxxx\n工作经历：...",
    height=200
)

# 获取内容
content = None
if uploaded_file:
    content = uploaded_file.read().decode("utf-8")
    st.success(f"已上传: {uploaded_file.name}")
elif resume_text:
    content = resume_text

# 美化选项
col1, col2 = st.columns(2)
with col1:
    style = st.selectbox("简历风格", ["专业", "简洁", "创意"])
with col2:
    target_job = st.text_input("目标岗位（可选）", placeholder="例如：Java开发工程师")

# 处理按钮
if st.button("🚀 美化并生成 DOC", type="primary"):
    if content:
        with st.spinner("AI 美化中..."):
            # AI 美化简历内容
            prompt = f"""你是一位专业的简历顾问。请美化以下简历内容，使其更加专业和有吸引力。

原始简历：
{content}

目标岗位：{target_job if target_job else "通用"}
风格：{style}

请输出美化后的简历内容，格式要求：
1. 保持原有信息完整
2. 使用更专业的表达
3. 突出亮点和成果
4. 量化工作成果（如有数字）

请按以下格式输出（用 === 分隔各部分）：
姓名 ===
联系方式 ===
求职意向 ===
个人优势 ===
工作经历 ===
项目经验 ===
教育背景 ===
技能证书 ===
"""
            
            response = client.chat.completions.create(
                model="qwen-turbo",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7
            )
            
            beautified_content = response.choices[0].message.content
        
        # 解析内容
        sections = {}
        current_section = None
        for line in beautified_content.split('\n'):
            if '===' in line:
                current_section = line.replace('===', '').strip()
                sections[current_section] = ""
            elif current_section:
                sections[current_section] += line + '\n'
        
        # 创建 DOC 文件
        doc = Document()
        
        # 设置默认字体
        style_doc = doc.styles['Normal']
        style_doc.font.name = 'Microsoft YaHei'
        style_doc.font.size = Pt(11)
        
        # 标题（姓名）
        if '姓名' in sections:
            title = doc.add_heading(sections['姓名'].strip(), level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 联系方式
        if '联系方式' in sections:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            runner = p.add_run(sections['联系方式'].strip())
            runner.font.size = Pt(10)
            runner.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()  # 空行
        
        # 各部分
        section_order = ['求职意向', '个人优势', '工作经历', '项目经验', '教育背景', '技能证书']
        
        for section_name in section_order:
            if section_name in sections and sections[section_name].strip():
                # 小标题
                heading = doc.add_heading(section_name, level=1)
                heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
                
                # 内容
                content_text = sections[section_name].strip()
                for line in content_text.split('\n'):
                    if line.strip():
                        p = doc.add_paragraph(line.strip(), style='List Bullet')
        
        # 保存到内存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # 显示预览
        st.subheader("📝 美化后的内容")
        st.text_area("预览", beautified_content, height=300)
        
        # 下载按钮
        st.download_button(
            label="📥 下载 DOC 文件",
            data=buffer,
            file_name="美化简历.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        st.success("✅ 简历美化完成！点击上方按钮下载 DOC 文件")
    else:
        st.warning("请上传简历文件或粘贴简历内容")

# 使用说明
with st.expander("📖 使用说明"):
    st.markdown("""
    ### 功能说明
    1. **上传简历**：支持 TXT 格式文件上传，或直接粘贴内容
    2. **AI 美化**：智能优化简历内容和表达
    3. **格式化输出**：自动生成排版美观的 DOC 文件
    
    ### 建议
    - 上传包含以下信息的简历效果更好：
      - 基本信息（姓名、联系方式）
      - 工作经历
      - 项目经验
      - 教育背景
      - 技能证书
    """)
