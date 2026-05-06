"""
生成不同格式的知识库测试文档
"""
import os

# 创建输出目录
output_dir = "E:\\newBeer\\study\\week2\\test_docs_multi"
os.makedirs(output_dir, exist_ok=True)

# ============================================
# 1. 生成 PDF 文档
# ============================================
print("生成 PDF 文档...")

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import cm

# 注册中文字体（使用 Windows 自带的 SimSun）
try:
    pdfmetrics.registerFont(TTFont('SimSun', 'C:\\Windows\\Fonts\\simsun.ttc', subfontIndex=0))
    font_name = 'SimSun'
except:
    # 如果字体加载失败，使用默认字体
    font_name = 'Helvetica'

pdf_file = f"{output_dir}\\考勤管理制度.pdf"
c = canvas.Canvas(pdf_file, pagesize=A4)
width, height = A4

# 标题
c.setFont(font_name, 18)
c.drawCentredString(width/2, height - 2*cm, "公司考勤管理制度")

# 内容
c.setFont(font_name, 12)
content_lines = [
    "",
    "一、工作时间",
    "1. 标准工作时间为每周一至周五",
    "2. 上班时间：上午 9:00",
    "3. 下班时间：下午 18:00",
    "4. 午休时间：12:00-13:00（1小时）",
    "",
    "二、打卡规定",
    "1. 员工每天需打卡4次：上班、午休、下午上班、下班",
    "2. 打卡方式：使用企业微信或考勤机",
    "3. 忘记打卡需在24小时内补卡，每月最多补卡3次",
    "",
    "三、迟到早退",
    "1. 迟到/早退 10分钟以内：口头警告",
    "2. 迟到/早退 10-30分钟：扣款50元",
    "3. 迟到/早退 30分钟以上：按旷工半天处理",
    "4. 月累计迟到3次以上：扣除当月绩效",
    "",
    "四、请假制度",
    "1. 请假需提前申请，经部门主管批准",
    "2. 病假需提供医院证明",
    "3. 年假需提前5个工作日申请",
    "",
    "五、加班管理",
    "1. 加班需提前申请，经部门主管批准",
    "2. 工作日加班按1.5倍工资计算",
    "3. 周末加班按2倍工资计算",
    "4. 法定节假日加班按3倍工资计算",
]

y_position = height - 3*cm
for line in content_lines:
    c.drawString(2*cm, y_position, line)
    y_position -= 0.6*cm

c.save()
print(f"✓ 已生成: {pdf_file}")

# ============================================
# 2. 生成 Word 文档
# ============================================
print("生成 Word 文档...")

from docx import Document

doc = Document()
doc.add_heading('公司培训管理制度', 0)

doc.add_heading('一、培训目标', level=1)
doc.add_paragraph('1. 提升员工专业技能和工作效率')
doc.add_paragraph('2. 增强团队协作能力')
doc.add_paragraph('3. 促进员工职业发展')

doc.add_heading('二、培训类型', level=1)
doc.add_paragraph('1. 新员工入职培训：为期3天，包含公司文化、规章制度、业务流程等')
doc.add_paragraph('2. 专业技能培训：定期组织，包含技术、管理、沟通等主题')
doc.add_paragraph('3. 外部培训：员工可申请参加外部培训课程，公司报销费用')
doc.add_paragraph('4. 在线学习：提供在线学习平台账号，鼓励自主学习')

doc.add_heading('三、培训申请流程', level=1)
doc.add_paragraph('1. 员工填写培训申请表')
doc.add_paragraph('2. 部门主管审批')
doc.add_paragraph('3. 人力资源部门审核')
doc.add_paragraph('4. 培训完成后提交培训总结')

doc.add_heading('四、培训费用', level=1)
doc.add_paragraph('1. 内部培训：免费')
doc.add_paragraph('2. 外部培训：单次不超过5000元，超过需总经理审批')
doc.add_paragraph('3. 在线学习平台：公司统一购买账号')

doc.add_heading('五、培训考核', level=1)
doc.add_paragraph('1. 培训后需进行考核或提交学习报告')
doc.add_paragraph('2. 考核结果纳入绩效考核')
doc.add_paragraph('3. 连续两次考核不合格者，取消当年培训申请资格')

doc.save(f"{output_dir}\\培训管理制度.docx")
print(f"✓ 已生成: {output_dir}\\培训管理制度.docx")

# ============================================
# 3. 生成 Excel 文档
# ============================================
print("生成 Excel 文档...")

import pandas as pd

# 员工福利表
data = {
    '福利项目': ['五险一金', '带薪年假', '节日福利', '生日福利', '健康体检', '交通补贴', '餐费补贴', '通讯补贴'],
    '适用对象': ['全体员工', '转正员工', '全体员工', '全体员工', '全体员工', '全体员工', '全体员工', '特定岗位'],
    '福利标准': [
        '按国家标准缴纳',
        '5-15天/年（根据工龄）',
        '春节、端午、中秋各500元',
        '生日当月发放200元购物卡',
        '每年一次，公司统一安排',
        '每月200元',
        '每月300元',
        '每月100-300元'
    ],
    '申请方式': [
        '入职自动办理',
        'OA系统申请',
        '自动发放',
        '自动发放',
        '人事部门通知',
        '自动发放',
        '自动发放',
        'OA系统申请'
    ],
    '备注': [
        '公积金比例12%',
        '需提前5天申请',
        '以购物卡形式发放',
        '以购物卡形式发放',
        '套餐标准500元/人',
        '无需发票',
        '无需发票',
        '需提供话费账单'
    ]
}

df = pd.DataFrame(data)
df.to_excel(f"{output_dir}\\员工福利表.xlsx", index=False)
print(f"✓ 已生成: {output_dir}\\员工福利表.xlsx")

# ============================================
# 4. 再生成一个部门职责表
# ============================================
dept_data = {
    '部门名称': ['技术部', '产品部', '运营部', '市场部', '人事部', '财务部'],
    '部门负责人': ['张工', '李经理', '王总', '赵经理', '陈主管', '刘会计'],
    '主要职责': [
        '负责产品技术研发、系统维护、技术支持',
        '负责产品规划、需求分析、产品设计',
        '负责用户运营、内容运营、活动策划',
        '负责市场推广、品牌建设、客户关系',
        '负责招聘、培训、绩效、薪酬福利',
        '负责财务核算、成本控制、税务管理'
    ],
    '人员编制': [15, 8, 12, 10, 5, 4],
    '联系电话': [
        '分机1001',
        '分机1002',
        '分机1003',
        '分机1004',
        '分机1005',
        '分机1006'
    ]
}

df_dept = pd.DataFrame(dept_data)
df_dept.to_excel(f"{output_dir}\\部门职责表.xlsx", index=False)
print(f"✓ 已生成: {output_dir}\\部门职责表.xlsx")

# ============================================
# 完成
# ============================================
print("\n" + "="*50)
print("✅ 所有测试文档生成完成！")
print(f"文件位置: {output_dir}")
print("\n生成的文件:")
print("1. 考勤管理制度.pdf")
print("2. 培训管理制度.docx")
print("3. 员工福利表.xlsx")
print("4. 部门职责表.xlsx")
